from enum import Enum
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Cm
from docx.oxml import parse_xml
from PIL import Image, ImageOps
from io import BytesIO


class ImageZoom(Enum):
    """
    图片缩放方式
    """
    FILL = 0  # 拉伸图像（不保持图像的比例），使其填充其包含元素。
    CONTAIN = 1  # 保持原始的宽高比，将图像缩放，使图像的长边能完全显示在容器内。这可能会在容器内部留下空白（填充）。
    COVER = 2  # 保持原始的宽高比，将图像缩放，使图像的短边能完全覆盖容器，长边可能会超出容器范围（裁剪）。


class ImageScale:
    """
    按照图片缩放方式,计算图片实际需要大小
    """

    def __init__(self, cell_width: int, cell_height: int,
                 zoom: ImageZoom = ImageZoom.FILL):
        """

        :param cell_width: 单元格宽 单位：厘米
        :param cell_height: 单元格高 单位：厘米
        :param zoom: 缩放方式
        """
        self.cell_width = cell_width
        self.cell_height = cell_height
        self.zoom = zoom

    def compute(self, image: Image) -> BytesIO:
        """
        计算图片的实际大小,返回
        :param image: 图片路径
        :return: BytesIO对象
        """
        img = image

        if self.zoom == ImageZoom.FILL:
            img = self.fill(img)
        elif self.zoom == ImageZoom.CONTAIN:
            img = self.contain(img)
        elif self.zoom == ImageZoom.COVER:
            img = self.cover(img)

        image_stream = BytesIO()
        if img.mode == 'RGBA':
            img = img.convert('RGB')
        img.save(image_stream, format='JPEG')
        # img.save(image_stream, format='JPEG')
        image_stream.seek(0)

        return image_stream

    def fill(self, img: Image) -> Image:
        """
        图片填充
        拉伸图像（不保持图像的比例），使其填充到单元格中
        :return: Image
        """
        # 计算目标长宽比
        return img

    def contain(self, img: Image) -> Image:
        """
        图片包含
        保持原始的宽高比，将图像缩放，使图像的长边能完全显示在容器内。这可能会在单元格内部留下空白。
        contain只会将图像的长宽比缩放到与容器的长宽比相同(短边方向填充)，而不会直接将图像大小设置为符合单元格的尺寸。
        这样做的目的是避免DPI未知转换cm与pixels的问题
        :return: Image
        """

        # 计算目标长宽比
        target_ratio = self.cell_width / self.cell_height

        # 计算图片的长宽比
        img_ratio = img.width / img.height

        # 如果图片的长宽比大于目标长宽比，那么调整图片的高度
        if img_ratio > target_ratio:
            new_height = int(round(img.width / target_ratio))
            img = ImageOps.pad(img, (img.width, new_height), color='white')

        # 如果图片的长宽比小于目标长宽比，那么调整图片的宽度
        elif img_ratio < target_ratio:
            new_width = int(round(img.height * target_ratio))
            img = ImageOps.pad(img, (new_width, img.height), color='white')

        return img

    def cover(self, img: Image) -> Image:
        """
        图片覆盖
        保持原始的宽高比，将图像缩放，使图像的短边能完全显示在容器内。这可能会在单元格外部裁剪图像。
        cover只会将图像的长宽比缩放到与容器的长宽比相同(长边方向裁剪)，而不会直接将图像大小设置为符合单元格的尺寸。
        这样做的目的是避免DPI未知转换cm与pixels的问题
        :return:
        """
        # 计算目标长宽比
        target_ratio = self.cell_width / self.cell_height

        # 计算图片的长宽比
        img_ratio = img.width / img.height

        # 如果图片的长宽比大于目标长宽比，那么裁剪图片的宽度
        if img_ratio > target_ratio:
            new_width = int(round(img.height * target_ratio))
            left = (img.width - new_width) / 2
            img = img.crop((left, 0, left + new_width, img.height))

        # 如果图片的长宽比小于目标长宽比，那么裁剪图片的高度
        elif img_ratio < target_ratio:
            new_height = int(round(img.width / target_ratio))
            top = (img.height - new_height) / 2
            img = img.crop((0, top, img.width, top + new_height))

        return img


class DocumentCreator:
    margin = [Cm(0.5), Cm(0.5), Cm(0.5), Cm(0.5)]  # 页边距，上下左右
    orientation = WD_ORIENT.PORTRAIT  # 文档方向 WD_ORIENT.PORTRAIT 纵向 WD_ORIENT.LANDSCAPE 横向
    size = [Cm(21), Cm(29.7)]

    def __init__(self, rows=1, cols=1, size=(21, 29.7), margin=(0.5, 0.5, 0.5, 0.5), zoom=ImageZoom.FILL):
        self.rows = rows  # 每一页的行数
        self.cols = cols  # 每一页的列数
        self.size = [Cm(size[0]), Cm(size[1])]  # 页面大小
        self.margin = [Cm(margin[0]), Cm(margin[1]), Cm(margin[2]), Cm(margin[3])]  # 页边距
        if size[0] > size[1]:
            self.orientation = WD_ORIENT.LANDSCAPE
        else:
            self.orientation = WD_ORIENT.PORTRAIT
        self.zoom = zoom  # 图片缩放方式

        self.document = Document()

    def setHorizontal(self, flag=True):
        """
        设置文档为横向，flag为True时
        false时设置为纵向
        :param flag:
        :return:
        """
        if flag:
            self.orientation = WD_ORIENT.LANDSCAPE
            self.size = [Inches(11.69), Inches(8.27)]
        else:
            self.orientation = WD_ORIENT.PORTRAIT
            self.size = [Inches(8.27), Inches(11.69)]

    def create(self, imageTempMemory):
        """
        创建文档
        :return:
        """

        self.setPageFormat()

        # 计算出每个图片的最大可能大小（减去两倍的边距）
        cell_width = Cm((self.size[0].cm - self.margin[2].cm - self.margin[3].cm) // self.cols)
        cell_height = Cm((self.size[1].cm - self.margin[0].cm - self.margin[1].cm) // self.rows)

        imageScale = ImageScale(cell_width.cm, cell_height.cm, self.zoom)

        table = None  # 初始化表格为None

        # 对于每个图片
        for index, picture in enumerate(imageTempMemory.images):
            i = index // (self.cols * self.rows)
            j = index % (self.cols * self.rows)

            # 如果是新的一页，创建一个新的表格
            if j == 0 and index != len(imageTempMemory.images):  # 检查是否还有更多的图片需要添加
                table = self.addTable(cell_width)

            # 计算行和列
            row = j // self.cols
            col = j % self.cols

            # 添加图片
            if index == len(imageTempMemory.images) - 1:
                picture = table.cell(row, col).paragraphs[0].add_run().add_picture(
                    imageScale.compute(picture), width=cell_width, height=Cm(cell_height.cm-0.2))  # 最后一张图片减去0.2厘米,防止出现空白页
            else:
                picture = table.cell(row, col).paragraphs[0].add_run().add_picture(
                    imageScale.compute(picture), width=cell_width, height=cell_height)

    def addTable(self, cell_width: int):
        table = self.document.add_table(self.rows, self.cols)
        table.autofit = False  # 关闭自动调整表格大小
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 设置表格水平居中
        for row in table.rows:
            for cell in row.cells:
                cell.width = cell_width  # 设置单元格宽度
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 设置单元格内容垂直居中
                # cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                # cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                cell._element.get_or_add_tcPr().append(parse_xml(
                    r'<w:noBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" />'))
                # 创建一个新的边距属性
                cell_margin = parse_xml(
                    r'<w:tcMar %s><w:top w:w="0" w:type="dxa"/><w:left w:w="0" w:type="dxa"/><w:bottom w:w="0" '
                    r'w:type="dxa"/><w:right w:w="0" w:type="dxa"/></w:tcMar>' % nsdecls(
                        'w'))

                # 设置单元格的边距
                cell._tc.get_or_add_tcPr().append(cell_margin)
        return table

    def setPageFormat(self):
        """
        设置文档页边距(统一设置所有节的页边距)
        :return:
        """
        # 获取第一个节
        for section in self.document.sections:
            # 设置页边距
            section.top_margin = self.margin[0]
            section.bottom_margin = self.margin[1]
            section.left_margin = self.margin[2]
            section.right_margin = self.margin[3]

            # 设置文档方向
            section.orientation = self.orientation

            # 设置文档的宽度和高度为A4纸的大小
            section.page_width = self.size[0]
            section.page_height = self.size[1]



if __name__ == '__main__':
    images = [r"D:\LUAO\Desktop\cstp\图片1.png", r"D:\LUAO\Desktop\cstp\图片2.png", r"D:\LUAO\Desktop\cstp\图片3.png"]
    doc = DocumentCreator(images, rows=1, cols=1, zoom=ImageZoom.COVER)
    doc.create()
    doc.document.save(r"test.docx")
